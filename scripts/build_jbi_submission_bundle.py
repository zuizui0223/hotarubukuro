#!/usr/bin/env python3
"""Safe entry point for the current JBI submission-bundle builder.

The core module owns document assembly. This entry point supplies the Markdown
inline renderer. Bold spans may contain scientific notation such as a*, L* and
C*, while the italic rule deliberately requires word boundaries so those
scientific suffixes are never mistaken for Markdown emphasis.

A small delivery-only normalization shortens the display labels in the wide
Appendix S6 candidate table and joins the final causal-ceiling sentence to the
preceding reproducibility paragraph. The underlying source values and meanings
are unchanged.
"""

from __future__ import annotations

import re

from docx.shared import Pt, RGBColor

import jbi_submission_bundle_core as core


TOKEN_PATTERN = re.compile(
    r"(\*\*[^\n]+?\*\*|`[^`\n]+`|\[[^\]]+\]\([^)]+\)|(?<![\w*])\*[^*\n]+?\*(?![\w*]))"
)
ORIGINAL_RENDER_MARKDOWN = core.render_markdown


def _style_run(run, *, bold: bool, italic: bool) -> None:
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def _append_inline(paragraph, text: str, *, bold: bool = False, italic: bool = False) -> None:
    position = 0
    for match in TOKEN_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            _style_run(run, bold=bold, italic=italic)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            _append_inline(paragraph, token[2:-2], bold=True, italic=italic)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            _style_run(run, bold=bold, italic=italic)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            _style_run(run, bold=bold, italic=True)
        else:
            label, url = core.normalize_markdown_link(token)
            run = paragraph.add_run(label)
            _style_run(run, bold=bold, italic=italic)
            if url:
                suffix = paragraph.add_run(f" ({url})")
                _style_run(suffix, bold=bold, italic=italic)
                suffix.font.color.rgb = RGBColor(60, 90, 130)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        _style_run(run, bold=bold, italic=italic)


def add_inline_runs(paragraph, text: str) -> None:
    _append_inline(paragraph, core.clean_markdown_text(text))


def normalize_delivery_markdown(text: str) -> str:
    if not text.startswith("# Appendix S6."):
        return text

    long_header = (
        "| Rank | Stable cell ID | Pigmented/observed | Predictive q | z | "
        "Neighbours/sites | Mean neighbour distance (km) | Mean environmental "
        "distance | 5-km population rank | Post-selection context class |"
    )
    short_header = (
        "| Rank | Cell ID | Pig./obs. | q | z | Neigh./sites | Mean dist. (km) | "
        "Mean env. dist. | Pop. rank (5 km) | Context class |"
    )
    text = text.replace(long_header, short_header)
    text = text.replace(
        "Human-context class is descriptive and was assigned only after candidate selection.",
        "Human-context class is descriptive and was assigned only after candidate selection. "
        "Table S6.4 abbreviates DID-proximate/high-population as `DID-high`, "
        "intermediate context as `Mid`, and remote/low-population as `Remote-low`.",
    )
    text = text.replace("DID-proximate, high population", "DID-high")
    text = text.replace("intermediate context", "Mid")
    text = text.replace("remote, low population", "Remote-low")
    text = text.replace(
        "It is retained as an execution verification rather than silently replacing the frozen numerical reference.\n\n"
        "The causal ceiling is unchanged:",
        "It is retained as an execution verification rather than silently replacing the frozen numerical reference. "
        "The causal ceiling is unchanged:",
    )
    return text


def render_markdown(document, text: str, **kwargs) -> None:
    ORIGINAL_RENDER_MARKDOWN(document, normalize_delivery_markdown(text), **kwargs)


core.add_inline_runs = add_inline_runs
core.render_markdown = render_markdown


if __name__ == "__main__":
    raise SystemExit(core.main())
