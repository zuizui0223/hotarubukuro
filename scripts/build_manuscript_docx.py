"""Build the submission-format DOCX from the locked manuscript Markdown.

The output is intentionally generated from the versioned Markdown rather than
from the obsolete Word draft. This keeps numerical claims auditable and makes
the Word document a reproducible publication artifact.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


DEFAULT_SOURCE = Path("manuscript/ecology-and-evolution-manuscript.md")
DEFAULT_OUTPUT = Path("local_outputs/manuscript_review/revised.docx")
FIGURE_FILES = {
    1: "figure_1_two_part_phenotype.png",
    2: "figure_2_environment_space.png",
    3: "figure_3_bombus_turnover.png",
    4: "figure_4_isolates_human_context.png",
}
FIGURE_ALT_TEXT = {
    1: (
        "Four-panel figure showing three common-extent maps of Japan, "
        "including Hokkaido: a map filled with each observation's extracted "
        "median sRGB colour, national white-like and pigmented records, and "
        "pigmented-only conditional intensity; the fourth panel shows the "
        "CIELAB a-star mixture boundary."
    ),
    2: (
        "Four-panel figure showing environmental posterior coefficients, "
        "SPDE spatial ranges, cross-fitted pigmentation probability, and "
        "cross-fitted pigmented-only intensity across Japan."
    ),
    3: (
        "Four-panel figure comparing paired held-out AUC with and without the "
        "predicted Bombus fingerprint and local turnover statistics against "
        "replicated natural-map reference distributions."
    ),
    4: (
        "Four-panel figure showing locally isolated pigmented cells, their "
        "natural-map extremeness checks, human-context contrasts, and early "
        "flowering and dark-colour predictive tail checks."
    ),
}


def set_run_font(run, name: str = "Times New Roman", size: float = 12) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def add_hyperlink(paragraph, text: str, url: str):
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    properties.append(size)
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


INLINE_TOKEN = re.compile(
    r"(\*\*.+?\*\*|(?<!\w)\*(?!\*)[^*\n]+?\*(?!\*)|`.+?`|https?://\S+)",
    flags=re.DOTALL,
)


def add_inline(paragraph, text: str, size: float = 12) -> None:
    cursor = 0
    for match in INLINE_TOKEN.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("http"):
            url = token.rstrip(".,;)")
            suffix = token[len(url) :]
            add_hyperlink(paragraph, url, url)
            if suffix:
                run = paragraph.add_run(suffix)
                set_run_font(run, size=size)
        else:
            run = paragraph.add_run(token.strip("*`"))
            set_run_font(run, size=size)
            if token.startswith("**"):
                run.bold = True
            elif token.startswith("*"):
                run.italic = True
            elif token.startswith("`"):
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                run.font.size = Pt(max(size - 1, 8))
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=10)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, result, end])


def set_picture_alt_text(run, description: str) -> None:
    drawing_properties = run._element.xpath(".//wp:docPr")
    if drawing_properties:
        drawing_properties[0].set("descr", description)
        drawing_properties[0].set("title", description)


def add_figure(document: Document, image_path: Path, alt_text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(15.8))
    set_picture_alt_text(run, alt_text)


def configure_section(section) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    line_numbers = OxmlElement("w:lnNumType")
    line_numbers.set(qn("w:countBy"), "1")
    line_numbers.set(qn("w:start"), "1")
    line_numbers.set(qn("w:restart"), "continuous")
    section._sectPr.append(line_numbers)
    add_page_number(section.footer.paragraphs[0])


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.widow_control = True

    for name, size, bold in (
        ("Title", 16, True),
        ("Heading 1", 14, True),
        ("Heading 2", 12, True),
        ("Heading 3", 12, True),
    ):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)
        if name != "Title":
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    caption = document.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.font.bold = False
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    caption.paragraph_format.space_after = Pt(6)


def parse_table(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if cells and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def shade_cell(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell_properties.append(shading)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    header_properties.append(header_marker)
    for row_index, row in enumerate(rows):
        for column_index in range(width):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_after = Pt(0)
            value = row[column_index] if column_index < len(row) else ""
            add_inline(paragraph, value, size=8.5)
            if row_index == 0:
                shade_cell(cell, "D9EAF7")
                for run in paragraph.runs:
                    run.bold = True
    document.add_paragraph()


def paragraph_blocks(lines: Iterable[str]):
    lines = list(lines)
    index = 0
    paragraph: list[str] = []
    while index < len(lines):
        line = lines[index].rstrip()
        if line.startswith("|"):
            if paragraph:
                yield ("paragraph", " ".join(paragraph))
                paragraph = []
            table_lines = []
            while index < len(lines) and lines[index].lstrip().startswith("|"):
                table_lines.append(lines[index].rstrip())
                index += 1
            yield ("table", table_lines)
            continue
        if not line.strip():
            if paragraph:
                yield ("paragraph", " ".join(paragraph))
                paragraph = []
        elif line.startswith("#"):
            if paragraph:
                yield ("paragraph", " ".join(paragraph))
                paragraph = []
            yield ("heading", line)
        elif re.match(r"^\d+\.\s+", line) or re.match(r"^-\s+", line):
            if paragraph:
                yield ("paragraph", " ".join(paragraph))
                paragraph = []
            yield ("list", line)
        else:
            paragraph.append(line.strip())
        index += 1
    if paragraph:
        yield ("paragraph", " ".join(paragraph))


def build(source: Path, output: Path) -> None:
    text = source.read_text(encoding="utf-8")
    document = Document()
    configure_styles(document)
    for section in document.sections:
        configure_section(section)

    title_seen = False
    abstract_seen = False
    references_seen = False
    figures_seen = 0
    for kind, payload in paragraph_blocks(text.splitlines()):
        if kind == "heading":
            hashes, heading = payload.split(" ", 1)
            level = len(hashes)
            if level == 1 and not title_seen:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(paragraph, heading, size=16)
                for run in paragraph.runs:
                    run.bold = True
                title_seen = True
                continue
            if heading == "Abstract" and not abstract_seen:
                document.add_page_break()
                abstract_seen = True
            references_seen = heading == "References" or references_seen
            style = "Heading 1" if level == 2 else "Heading 2"
            paragraph = document.add_paragraph(style=style)
            add_inline(paragraph, heading, size=14 if level == 2 else 12)
            continue

        if kind == "table":
            add_table(document, parse_table(payload))
            continue

        if kind == "list":
            list_text = re.sub(r"^(?:\d+\.|-)\s+", "", payload)
            style = (
                "List Number"
                if re.match(r"^\d+\.\s+", payload)
                else "List Bullet"
            )
            paragraph = document.add_paragraph(style=style)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            add_inline(paragraph, list_text)
            continue

        figure_match = re.match(r"\*\*Figure ([1-4])\.", payload)
        if figure_match:
            figure_number = int(figure_match.group(1))
            image_path = source.parent / "figures" / FIGURE_FILES[figure_number]
            if not image_path.exists():
                raise FileNotFoundError(
                    f"Missing Figure {figure_number} image: {image_path}"
                )
            if figures_seen:
                document.add_page_break()
            add_figure(
                document,
                image_path,
                FIGURE_ALT_TEXT[figure_number],
            )
            figures_seen += 1
            paragraph = document.add_paragraph(style="Caption")
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_after = Pt(6)
            add_inline(paragraph, payload, size=10)
            continue

        paragraph = document.add_paragraph()
        if not abstract_seen:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_after = Pt(6)
        elif references_seen:
            paragraph.paragraph_format.left_indent = Cm(0.75)
            paragraph.paragraph_format.first_line_indent = Cm(-0.75)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_after = Pt(6)
        add_inline(paragraph, payload)

    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    core = document.core_properties
    core.title = (
        "From hiking records to two-part floral-colour geography: "
        "scale-aware inference from YAMAP photographs of Campanula punctata"
    )
    core.author = "Ruiqi Zhang"
    core.subject = "Ecology and Evolution Research Article"
    core.keywords = (
        "bumblebee community; digital phenotyping; flower colour; "
        "INLA-SPDE; intraspecific trait variation; YAMAP"
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"Manuscript DOCX written: {output.resolve()}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    arguments = parser.parse_args()
    build(arguments.source, arguments.output)


if __name__ == "__main__":
    main()
