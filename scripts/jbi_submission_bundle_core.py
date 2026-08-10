#!/usr/bin/env python3
"""Build the current JBI review/submission bundle from versioned sources.

The builder keeps the double-anonymous scientific package separate from
identifying or still-unresolved portal materials. It creates editable DOCX
files, embeds the four validated Main figures, combines Appendices S1-S6 into
one Supporting Information file, and writes a readiness report rather than
inventing author, funding, repository-link, disclosure or taxon-image details.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import re
import shutil
import sys
import zipfile
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable, Iterator, Mapping, Sequence

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TITLE = "From broad geography to local boundaries: flower-colour biogeography from hiking photographs"
RUNNING_TITLE = "Layered flower-colour biogeography"

MAIN_SOURCE = Path("submission/jbi/JBI_main_manuscript_anonymized.md")
TITLE_PAGE_SOURCE = Path("submission/jbi/JBI_title_page_template.md")
COVER_SOURCE = Path("submission/jbi/JBI_cover_letter.md")
TRANSLATED_ABSTRACT_SOURCE = Path("submission/jbi/JBI_translated_abstract_ja.md")
CHECKLIST_SOURCE = Path("submission/jbi/JBI_submission_checklist.md")
CAPTIONS_SOURCE = Path("submission/jbi/JBI_main_figure_captions.md")
SDM_CHECKLIST_SOURCE = Path("submission/jbi/JBI_sdm_model_building_checklist.md")
SUPPORTING_SOURCES = tuple(
    Path(f"submission/jbi/supporting/Appendix_S{i}_{name}.md")
    for i, name in (
        (1, "yamap_public_benchmark"),
        (2, "image_phenotyping"),
        (3, "broad_environment_spatial_model"),
        (4, "bombus_sdm_occurrence_support"),
        (5, "local_pollinator_robustness"),
        (6, "event_departures_human_context"),
    )
)

FIGURE_STEMS = (
    "Figure_1_measurement_two_part_phenotype",
    "Figure_2_broad_environment_spatial_template",
    "Figure_3_local_focal_bombus_boundaries",
    "Figure_4_calibrated_local_departures",
)

FORBIDDEN_IDENTIFIERS = (
    "zuizui0223",
    "rachelzhang",
    "ZHANG Ruiqi",
    "張瑞琪",
)

PLACEHOLDER_PATTERNS = {
    "authors_and_orcids": ("[Author 1 full name]", "[ORCID]"),
    "affiliations": ("[Department / Graduate School", "[Additional affiliation]"),
    "corresponding_author": ("[One corresponding author only]", "[email]"),
    "acknowledgements": ("[Insert acknowledgements",),
    "funding": ("[Insert funding sources",),
    "conflict_of_interest": ("[Insert agreed statement",),
    "author_contributions": ("Conceptualization: [ ]",),
    "cover_letter_signature": ("[Corresponding author]", "[Affiliation]", "[Email]"),
}


@dataclass(frozen=True)
class BundleFile:
    role: str
    path: Path
    required_for_review: bool
    required_for_portal: bool


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(4 * 1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def write_text_atomic(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(f".{path.name}.tmp")
    temporary.write_text(text, encoding="utf-8")
    temporary.replace(path)


def write_json_atomic(path: Path, value: Mapping[str, object]) -> None:
    write_text_atomic(
        path,
        json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
    )


def ensure_style(document: DocumentType, name: str, base: str | None = None):
    styles = document.styles
    if name in styles:
        return styles[name]
    style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if base and base in styles:
        style.base_style = styles[base]
    return style


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def configure_document(document: DocumentType, *, anonymized: bool) -> None:
    properties = document.core_properties
    properties.title = TITLE
    properties.subject = "Journal of Biogeography submission material"
    properties.author = "" if anonymized else "Author details pending"
    properties.last_modified_by = ""
    properties.comments = "Generated from versioned manuscript sources."
    properties.keywords = "biogeography; flower colour; iEcology; pollination"

    for section in document.sections:
        section.top_margin = Inches(0.72)
        section.bottom_margin = Inches(0.72)
        section.left_margin = Inches(0.78)
        section.right_margin = Inches(0.78)
        section.header_distance = Inches(0.30)
        section.footer_distance = Inches(0.30)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True

    title_style = document.styles["Title"]
    title_style.font.name = "Arial"
    title_style.font.size = Pt(16)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(37, 49, 59)
    title_style.paragraph_format.space_after = Pt(10)

    for index, size in ((1, 14), (2, 12), (3, 10.8), (4, 10.2)):
        style_name = f"Heading {index}"
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(37, 49, 59)
        style.paragraph_format.space_before = Pt(10 if index <= 2 else 7)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True

    caption = ensure_style(document, "Figure Caption", "Normal")
    caption.font.name = "Arial"
    caption.font.size = Pt(9.5)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(7)
    caption.paragraph_format.keep_with_next = True

    code = ensure_style(document, "Code Block", "Normal")
    code.font.name = "Courier New"
    code.font.size = Pt(8.5)
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(5)

    quote = ensure_style(document, "Source Note", "Normal")
    quote.font.italic = True
    quote.font.color.rgb = RGBColor(88, 99, 110)
    quote.paragraph_format.left_indent = Inches(0.28)
    quote.paragraph_format.right_indent = Inches(0.18)


def clean_markdown_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = text.replace("<br>", "\n").replace("<br/>", "\n")
    return text


def markdown_cells(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [cell.strip().replace("\\|", "|") for cell in stripped.split("|")]


def is_table_separator(line: str) -> bool:
    cells = markdown_cells(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def normalize_markdown_link(token: str) -> tuple[str, str | None]:
    match = re.fullmatch(r"\[([^]]+)]\(([^)]+)\)", token)
    if not match:
        return token, None
    return match.group(1), match.group(2)


def add_inline_runs(paragraph, text: str) -> None:
    text = clean_markdown_text(text)
    token_pattern = re.compile(
        r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^)]+\))"
    )
    position = 0
    for match in token_pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            label, url = normalize_markdown_link(token)
            paragraph.add_run(label)
            if url:
                run = paragraph.add_run(f" ({url})")
                run.font.color.rgb = RGBColor(60, 90, 130)
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def add_markdown_table(document: DocumentType, lines: Sequence[str]) -> None:
    header = markdown_cells(lines[0])
    rows = [markdown_cells(line) for line in lines[2:]]
    column_count = len(header)
    if column_count == 0:
        return
    table = document.add_table(rows=1, cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    for index, value in enumerate(header):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E9EEF2")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_inline_runs(paragraph, value)
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(8.5)
    for raw_row in rows:
        values = raw_row + [""] * max(0, column_count - len(raw_row))
        cells = table.add_row().cells
        for index in range(column_count):
            cell = cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_runs(paragraph, values[index] if index < len(values) else "")
            for run in paragraph.runs:
                run.font.size = Pt(8.3)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def next_special(lines: Sequence[str], index: int) -> bool:
    if index >= len(lines):
        return True
    line = lines[index]
    stripped = line.strip()
    if not stripped:
        return True
    if stripped.startswith("```"):
        return True
    if re.match(r"^#{1,6}\s+", stripped):
        return True
    if re.match(r"^(?:[-*+] |\d+\. )", stripped):
        return True
    if stripped.startswith(">"):
        return True
    if index + 1 < len(lines) and "|" in line and is_table_separator(lines[index + 1]):
        return True
    return False


def render_markdown(
    document: DocumentType,
    text: str,
    *,
    first_heading_as_title: bool = False,
    page_break_before_first_heading: bool = False,
) -> None:
    lines = clean_markdown_text(text).splitlines()
    index = 0
    heading_count = 0
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            paragraph = document.add_paragraph(style="Code Block")
            if language:
                label = paragraph.add_run(f"[{language}]\n")
                label.bold = True
            paragraph.add_run("\n".join(code_lines))
            continue

        if index + 1 < len(lines) and "|" in raw and is_table_separator(lines[index + 1]):
            table_lines = [raw, lines[index + 1]]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(document, table_lines)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            heading_count += 1
            if page_break_before_first_heading and heading_count == 1:
                document.add_page_break()
            if first_heading_as_title and heading_count == 1 and level == 1:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_runs(paragraph, heading_text)
            else:
                use_level = min(max(level, 1), 4)
                paragraph = document.add_paragraph(style=f"Heading {use_level}")
                add_inline_runs(paragraph, heading_text)
                keep_with_next(paragraph)
            index += 1
            continue

        list_match = re.match(r"^\s*([-*+]|\d+\.)\s+(.+)$", raw)
        if list_match:
            marker, value = list_match.groups()
            style = "List Number" if marker[0].isdigit() else "List Bullet"
            paragraph = document.add_paragraph(style=style)
            paragraph.paragraph_format.space_after = Pt(2)
            add_inline_runs(paragraph, value)
            index += 1
            continue

        if stripped.startswith(">"):
            paragraph = document.add_paragraph(style="Source Note")
            add_inline_runs(paragraph, stripped.lstrip("> "))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines) and not next_special(lines, index):
            paragraph_lines.append(lines[index].strip())
            index += 1
        paragraph = document.add_paragraph()
        add_inline_runs(paragraph, " ".join(paragraph_lines))


def parse_figure_captions(text: str) -> list[tuple[int, str, str]]:
    lines = text.splitlines()
    captures: list[tuple[int, str, str]] = []
    current_number: int | None = None
    current_title = ""
    body: list[str] = []
    for line in lines:
        match = re.match(r"^##\s+Figure\s+([1-4])\.\s*(.*)$", line.strip())
        if match:
            if current_number is not None:
                captures.append((current_number, current_title, " ".join(body).strip()))
            current_number = int(match.group(1))
            current_title = match.group(2).strip()
            body = []
        elif current_number is not None and line.strip():
            body.append(line.strip())
    if current_number is not None:
        captures.append((current_number, current_title, " ".join(body).strip()))
    if [number for number, _, _ in captures] != [1, 2, 3, 4]:
        raise ValueError("Figure captions must contain Figures 1-4 exactly once and in order")
    return captures


def add_embedded_figures(
    document: DocumentType,
    captions_text: str,
    figure_root: Path,
) -> None:
    document.add_page_break()
    heading = document.add_paragraph(style="Heading 1")
    add_inline_runs(heading, "Figure legends and embedded figures")
    captions = parse_figure_captions(captions_text)
    for number, title, body in captions:
        if number > 1:
            document.add_page_break()
        paragraph = document.add_paragraph(style="Figure Caption")
        label = paragraph.add_run(f"Figure {number}. ")
        label.bold = True
        if title:
            run = paragraph.add_run(title)
            run.bold = True
            if body:
                paragraph.add_run(" ")
        if body:
            add_inline_runs(paragraph, body)
        image = figure_root / f"{FIGURE_STEMS[number - 1]}.png"
        if not image.is_file():
            raise FileNotFoundError(f"Missing Main figure: {image}")
        picture_paragraph = document.add_paragraph()
        picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture_paragraph.add_run().add_picture(str(image), width=Inches(6.68))


def save_document(document: DocumentType, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)


def build_main_document(output: Path, figure_root: Path) -> None:
    document = Document()
    configure_document(document, anonymized=True)
    render_markdown(
        document,
        MAIN_SOURCE.read_text(encoding="utf-8"),
        first_heading_as_title=True,
    )
    add_embedded_figures(
        document,
        CAPTIONS_SOURCE.read_text(encoding="utf-8"),
        figure_root,
    )
    save_document(document, output)


def build_simple_document(
    source: Path,
    output: Path,
    *,
    anonymized: bool,
    first_heading_as_title: bool = True,
) -> None:
    document = Document()
    configure_document(document, anonymized=anonymized)
    render_markdown(
        document,
        source.read_text(encoding="utf-8"),
        first_heading_as_title=first_heading_as_title,
    )
    save_document(document, output)


def build_supporting_document(output: Path) -> None:
    document = Document()
    configure_document(document, anonymized=True)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Supporting Information")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(TITLE)
    subtitle_run.italic = True
    note = document.add_paragraph(style="Source Note")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        "Combined editable file assembled from the six current manuscript-facing appendices."
    )
    for index, source in enumerate(SUPPORTING_SOURCES):
        if not source.is_file():
            raise FileNotFoundError(f"Missing Supporting Information source: {source}")
        render_markdown(
            document,
            source.read_text(encoding="utf-8"),
            first_heading_as_title=False,
            page_break_before_first_heading=True,
        )
    save_document(document, output)


def copy_figures(figure_root: Path, destination: Path) -> list[Path]:
    destination.mkdir(parents=True, exist_ok=True)
    copied: list[Path] = []
    for stem in FIGURE_STEMS:
        for suffix in (".png", ".pdf"):
            source = figure_root / f"{stem}{suffix}"
            if not source.is_file():
                raise FileNotFoundError(f"Missing generated figure: {source}")
            target = destination / source.name
            shutil.copy2(source, target)
            copied.append(target)
    return copied


def document_xml_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        chunks: list[str] = []
        for name in archive.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                chunks.append(archive.read(name).decode("utf-8", errors="replace"))
        return "\n".join(chunks)


def detect_readiness(
    output_dir: Path,
    title_text: str,
    cover_text: str,
    main_text: str,
) -> dict[str, object]:
    blockers: list[dict[str, str]] = []
    warnings: list[dict[str, str]] = []

    for key, tokens in PLACEHOLDER_PATTERNS.items():
        haystack = cover_text if key == "cover_letter_signature" else title_text
        if any(token in haystack for token in tokens):
            blockers.append(
                {
                    "id": key,
                    "status": "pending_author_input",
                    "detail": "Identifying/co-author-approved content remains a template placeholder.",
                }
            )

    taxon_candidates = tuple(Path("submission/jbi").glob("taxon_image.*"))
    if not taxon_candidates:
        blockers.append(
            {
                "id": "taxon_image",
                "status": "missing",
                "detail": "JBI requires a separate study-taxon image with ownership/permission documented.",
            }
        )

    repository_markers = (
        "PRIVATE_REVIEW_URL",
        "private Dryad URL",
        "private-for-peer-review",
        "[review repository",
    )
    if not any(marker.casefold() in main_text.casefold() for marker in repository_markers):
        blockers.append(
            {
                "id": "private_review_data_code_link",
                "status": "missing",
                "detail": "Insert the randomized private-for-peer-review Dryad/equivalent URL before portal upload.",
            }
        )

    warnings.extend(
        [
            {
                "id": "ai_use_disclosure_decision",
                "status": "author_decision_required",
                "detail": (
                    "Confirm the Wiley/JBI disclosure required for any generative-AI use in manuscript development; "
                    "routine spelling/grammar editing is treated differently by the policy."
                ),
            },
            {
                "id": "map_scale_and_projection_final_check",
                "status": "editorial_check_required",
                "detail": (
                    "The current maps display latitude/longitude. Confirm final captions/proofs state the projection "
                    "and supply scale information as required by the current JBI map guidance."
                ),
            },
            {
                "id": "portal_field_designations",
                "status": "final_portal_check_required",
                "detail": "Re-check the live Wiley Authors portal file-type labels immediately before upload.",
            },
        ]
    )

    main_doc = output_dir / "01_Main_Manuscript_Anonymized.docx"
    anonymous_xml = document_xml_text(main_doc)
    leaked = [token for token in FORBIDDEN_IDENTIFIERS if token.casefold() in anonymous_xml.casefold()]
    if leaked:
        blockers.append(
            {
                "id": "anonymity_failure",
                "status": "failed",
                "detail": "Identifying strings found in anonymized DOCX: " + ", ".join(leaked),
            }
        )

    return {
        "schema_version": 1,
        "generated_utc": datetime.now(timezone.utc).isoformat(),
        "review_science_bundle_complete": not any(
            blocker["id"] == "anonymity_failure" for blocker in blockers
        ),
        "portal_ready": len(blockers) == 0,
        "blockers": blockers,
        "warnings": warnings,
        "interpretation": (
            "The anonymous scientific package is assembled and structurally checked. "
            "Portal readiness remains false until author-controlled identifying, disclosure, repository-link and taxon-image items are completed."
        ),
    }


def render_readiness_markdown(readiness: Mapping[str, object]) -> str:
    lines = [
        "# JBI submission-bundle readiness",
        "",
        f"- Anonymous scientific bundle complete: **{str(readiness['review_science_bundle_complete']).upper()}**",
        f"- Portal ready: **{str(readiness['portal_ready']).upper()}**",
        "",
        str(readiness["interpretation"]),
        "",
        "## Portal blockers",
        "",
    ]
    blockers = readiness.get("blockers", [])
    if blockers:
        for item in blockers:  # type: ignore[assignment]
            lines.append(f"- **{item['id']}** — {item['status']}: {item['detail']}")
    else:
        lines.append("- None.")
    lines.extend(["", "## Required final checks", ""])
    for item in readiness.get("warnings", []):  # type: ignore[assignment]
        lines.append(f"- **{item['id']}** — {item['status']}: {item['detail']}")
    lines.extend(
        [
            "",
            "## Current assembled files",
            "",
            "- anonymized Main DOCX with four embedded figures and figure legends;",
            "- one combined editable Supporting Information DOCX containing Appendices S1-S6;",
            "- separate title-page, cover-letter and translated-abstract DOCX templates;",
            "- separate SDM/model-building checklist DOCX;",
            "- four 600-dpi PNG and four vector PDF figure files;",
            "- file/hashing manifest and this readiness report.",
            "",
        ]
    )
    return "\n".join(lines)


def write_inventory(files: Sequence[BundleFile], root: Path, output: Path) -> None:
    rows = []
    for item in files:
        relative = item.path.relative_to(root)
        rows.append(
            {
                "role": item.role,
                "path": relative.as_posix(),
                "required_for_review": item.required_for_review,
                "required_for_portal": item.required_for_portal,
                "size_bytes": item.path.stat().st_size,
                "sha256": sha256_file(item.path),
            }
        )
    output.parent.mkdir(parents=True, exist_ok=True)
    with output.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)


def write_bundle_readme(path: Path) -> None:
    text = f"""# JBI review/submission bundle

Article: {TITLE}

This bundle separates a complete anonymous scientific review package from
portal materials that still require author/co-author input.

## Upload-oriented files

- `01_Main_Manuscript_Anonymized.docx`: Main text, legends and embedded Figures 1-4.
- `02_Supporting_Information_Appendices_S1-S6.docx`: one combined editable SI file.
- `03_Title_Page_TEMPLATE.docx`: identifying fields; complete before upload.
- `04_Cover_Letter_TEMPLATE.docx`: sign-off fields; complete before upload.
- `05_Translated_Abstract_Japanese.docx`: optional translated abstract.
- `06_JBI_SDM_Model_Building_Checklist.docx`: current Bombus SDM reporting checklist.
- `Figures/`: separate 600-dpi PNG and vector PDF files.

## Readiness

Read `Submission_Readiness.md` before uploading anything. `portal_ready=false` is
expected until authorship, affiliations, ORCIDs, acknowledgements, funding,
conflict of interest, author contributions, private review repository URL and
taxon image are finalized by the authors.

The underlying manuscript source, code and frozen evidence remain versioned in
the repository; this directory is a generated delivery bundle.
"""
    write_text_atomic(path, text)


def make_zip(source_dir: Path, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(destination, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=6) as archive:
        for path in sorted(source_dir.rglob("*")):
            if not path.is_file() or path == destination:
                continue
            archive.write(path, path.relative_to(source_dir).as_posix())


def validate_sources(figure_root: Path) -> None:
    sources = (
        MAIN_SOURCE,
        TITLE_PAGE_SOURCE,
        COVER_SOURCE,
        TRANSLATED_ABSTRACT_SOURCE,
        CHECKLIST_SOURCE,
        CAPTIONS_SOURCE,
        SDM_CHECKLIST_SOURCE,
        *SUPPORTING_SOURCES,
    )
    missing = [str(path) for path in sources if not path.is_file()]
    for stem in FIGURE_STEMS:
        for suffix in (".png", ".pdf"):
            candidate = figure_root / f"{stem}{suffix}"
            if not candidate.is_file():
                missing.append(str(candidate))
    if missing:
        raise FileNotFoundError("Missing bundle inputs: " + ", ".join(missing))


def build_bundle(figure_root: Path, output_dir: Path) -> dict[str, object]:
    validate_sources(figure_root)
    if output_dir.exists():
        shutil.rmtree(output_dir)
    output_dir.mkdir(parents=True)

    main_doc = output_dir / "01_Main_Manuscript_Anonymized.docx"
    supporting_doc = output_dir / "02_Supporting_Information_Appendices_S1-S6.docx"
    title_doc = output_dir / "03_Title_Page_TEMPLATE.docx"
    cover_doc = output_dir / "04_Cover_Letter_TEMPLATE.docx"
    translated_doc = output_dir / "05_Translated_Abstract_Japanese.docx"
    sdm_doc = output_dir / "06_JBI_SDM_Model_Building_Checklist.docx"

    build_main_document(main_doc, figure_root)
    build_supporting_document(supporting_doc)
    build_simple_document(TITLE_PAGE_SOURCE, title_doc, anonymized=False)
    build_simple_document(COVER_SOURCE, cover_doc, anonymized=False)
    build_simple_document(TRANSLATED_ABSTRACT_SOURCE, translated_doc, anonymized=True)
    build_simple_document(SDM_CHECKLIST_SOURCE, sdm_doc, anonymized=True)

    figure_files = copy_figures(figure_root, output_dir / "Figures")
    shutil.copy2(CHECKLIST_SOURCE, output_dir / "JBI_Submission_Checklist.md")
    write_bundle_readme(output_dir / "README.md")
    write_text_atomic(
        output_dir / "Taxon_Image_REQUIRED.txt",
        "JBI requires a separate high-quality taxon image. Add an author-owned or permission-cleared image before portal upload and document ownership/permission.\n",
    )
    write_text_atomic(
        output_dir / "Private_Review_Data_Code_Link_REQUIRED.txt",
        "Insert the randomized private-for-peer-review Dryad or equivalent URL in the manuscript/Supplementary Files before portal upload. Do not substitute an author-identifying public repository in the anonymized review file.\n",
    )

    readiness = detect_readiness(
        output_dir,
        TITLE_PAGE_SOURCE.read_text(encoding="utf-8"),
        COVER_SOURCE.read_text(encoding="utf-8"),
        MAIN_SOURCE.read_text(encoding="utf-8"),
    )
    write_json_atomic(output_dir / "Submission_Readiness.json", readiness)
    write_text_atomic(
        output_dir / "Submission_Readiness.md",
        render_readiness_markdown(readiness),
    )

    bundle_files = [
        BundleFile("main_anonymized", main_doc, True, True),
        BundleFile("supporting_information", supporting_doc, True, True),
        BundleFile("title_page_template", title_doc, False, True),
        BundleFile("cover_letter_template", cover_doc, False, True),
        BundleFile("translated_abstract", translated_doc, False, False),
        BundleFile("sdm_model_building_checklist", sdm_doc, True, True),
        BundleFile("submission_checklist", output_dir / "JBI_Submission_Checklist.md", False, False),
        BundleFile("readiness_markdown", output_dir / "Submission_Readiness.md", True, True),
        BundleFile("readiness_json", output_dir / "Submission_Readiness.json", True, True),
        BundleFile("taxon_image_placeholder", output_dir / "Taxon_Image_REQUIRED.txt", False, True),
        BundleFile("private_review_link_placeholder", output_dir / "Private_Review_Data_Code_Link_REQUIRED.txt", True, True),
        BundleFile("bundle_readme", output_dir / "README.md", True, False),
    ]
    for path in figure_files:
        bundle_files.append(BundleFile("main_figure", path, True, True))

    manifest = output_dir / "Submission_File_Manifest.csv"
    write_inventory(bundle_files, output_dir, manifest)
    bundle_files.append(BundleFile("file_manifest", manifest, True, True))

    archive = output_dir.parent / "JBI_Review_Submission_Bundle.zip"
    if archive.exists():
        archive.unlink()
    make_zip(output_dir, archive)

    summary = {
        "schema_version": 1,
        "output_directory": output_dir.as_posix(),
        "archive": archive.as_posix(),
        "archive_sha256": sha256_file(archive),
        "archive_size_bytes": archive.stat().st_size,
        "main_figures": 4,
        "supporting_appendices": 6,
        "portal_ready": readiness["portal_ready"],
        "review_science_bundle_complete": readiness["review_science_bundle_complete"],
        "files_in_manifest": len(bundle_files),
    }
    write_json_atomic(output_dir.parent / "JBI_Submission_Bundle_Summary.json", summary)
    return summary


def parser() -> argparse.ArgumentParser:
    value = argparse.ArgumentParser(description=__doc__)
    value.add_argument(
        "--figure-root",
        type=Path,
        default=Path("results/jbi_figure_bundle/main_figures"),
        help="Directory containing the four validated PNG/PDF pairs.",
    )
    value.add_argument(
        "--output",
        type=Path,
        default=Path("results/jbi_submission_bundle/package"),
        help="Generated package directory.",
    )
    return value


def main(argv: Sequence[str] | None = None) -> int:
    arguments = parser().parse_args(argv)
    summary = build_bundle(arguments.figure_root, arguments.output)
    print(json.dumps(summary, ensure_ascii=False, indent=2, sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
